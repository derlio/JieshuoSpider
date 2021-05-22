import os

from docx import Document


def main():
    input_dir = 'files/output/docs'
    output_dir = 'files/output/docs2'

    # 检查输出目录是否存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 遍历所有原文件
    for file in os.listdir(input_dir):
        try:
            print('transforming: ' + file)
            doc = Document(os.path.join(input_dir, file))
            template = Document('template/template.docx')
            idx = 0
            for para in doc.paragraphs:
                # 如果para.text是空或者有关键字，不添加到新文件中，这样就把二维码和扫码添加的
                # 的文字去除了
                if para.text is None or '扫码' in para.text:
                    continue

                if idx == 0:
                    # 第一行默认为标题
                    template.add_heading(para.text, level=2)
                else:
                    template.add_paragraph(para.text)
                idx += 1

            template.save(os.path.join(output_dir, file))
        except Exception as err:
            print(err)


if __name__ == '__main__':
    main()
